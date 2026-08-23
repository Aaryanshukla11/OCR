import io
import os
from typing import List, Tuple, Optional
import numpy as np
from PIL import Image, ImageDraw, ImageFont

class DocLoader:
    """
    Input layer document loader for Word documents (.doc and .docx).
    Extracts structured text, paragraphs, tables, and renders document pages to image arrays.
    """
    @staticmethod
    def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
        text_parts = []
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
                    
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        text_parts.append(" | ".join(row_cells))
        except Exception:
            # Fallback text extraction if docx library is not available or file is plain text/raw doc
            try:
                raw_str = file_bytes.decode('utf-8', errors='ignore')
                lines = [line.strip() for line in raw_str.splitlines() if line.strip()]
                text_parts = lines[:50]
            except Exception:
                text_parts = ["Word Document Content"]

        return "\n".join(text_parts) if text_parts else "Empty Document"

    @staticmethod
    def render_doc_pages(file_bytes: bytes, filename: str) -> List[Tuple[int, np.ndarray, int, int, str]]:
        """
        Renders Word document content to structured page image representations.
        Returns a list of tuples: [(page_number, np_array_img, width, height, extracted_text), ...]
        """
        extracted_text = DocLoader.extract_text_from_docx_bytes(file_bytes)
        lines = [line for line in extracted_text.splitlines() if line.strip()]
        
        # Paginate lines (~25 lines per page)
        lines_per_page = 25
        pages_text_chunks = []
        for i in range(0, max(1, len(lines)), lines_per_page):
            chunk = lines[i:i + lines_per_page]
            pages_text_chunks.append("\n".join(chunk))

        page_results = []
        width, height = 800, 1050
        
        for idx, page_text in enumerate(pages_text_chunks):
            page_num = idx + 1
            # Create a clean document page canvas (white background)
            page_img = Image.new("RGB", (width, height), color=(255, 255, 255))
            draw = ImageDraw.Draw(page_img)
            
            # Header line
            draw.rectangle([40, 30, width - 40, 32], fill=(200, 200, 200))
            draw.text((40, 10), f"{filename} - Page {page_num}", fill=(100, 100, 100))
            
            y_offset = 60
            page_lines = page_text.splitlines()
            for line in page_lines:
                draw.text((50, y_offset), line, fill=(20, 20, 20))
                y_offset += 35
                if y_offset > height - 60:
                    break
                    
            page_results.append((page_num, np.array(page_img), width, height, page_text))
            
        return page_results
